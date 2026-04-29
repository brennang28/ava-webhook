import copy
from docx import Document
from docx.shared import RGBColor
import re

template_path = 'assets/Aschettino, Ava - Cover Letter Template.docx'
doc = Document(template_path)

salutation_idx = -1
for i, p in enumerate(doc.paragraphs):
    if "Dear" in p.text:
        salutation_idx = i
        break

body_para = doc.paragraphs[salutation_idx + 1]

# Clone body_para XML
new_p = copy.deepcopy(body_para._element)
# Clear runs
new_p.clear_content()

# add it
doc.paragraphs[salutation_idx]._element.addnext(new_p)

# wrap in Paragraph object
from docx.text.paragraph import Paragraph
p = Paragraph(new_p, doc.paragraphs[salutation_idx]._parent)
p.add_run("This is cloned paragraph.")

doc.save('scratch/cloned.docx')
