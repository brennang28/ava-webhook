import copy
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

def _is_header_line(line, company_name):
    header_patterns = [
        r"^Ava Aschettino$",
        r"^New York, NY$",
        r"^\(516\) 532-3384$",
        r"^avaaschettino@gmail\.com$",
        r"^Hiring Manager$",
        r"^" + re.escape(company_name) + r"$",
        r"^\[Company Address\]$",
        r"^Sincerely,$",
        r"^Best regards,$",
        r"^Thank you$",
    ]
    stripped = line.strip()
    for pattern in header_patterns:
        if re.search(pattern, stripped, re.IGNORECASE):
            return True
    if re.match(r'^Dear\s+', stripped, re.IGNORECASE):
        return True
    return False

def test_cloning():
    doc = Document('assets/Aschettino, Ava - Cover Letter Template.docx')
    company_name = "Lambent"
    
    salutation_idx = -1
    closing_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "Dear" in p.text:
            salutation_idx = i
        if "Sincerely" in p.text:
            closing_idx = i
            
    closing_texts = []
    if closing_idx != -1 and closing_idx > salutation_idx:
        for i in range(closing_idx, len(doc.paragraphs)):
            text = doc.paragraphs[i].text
            if text.strip():
                closing_texts.append(text)

    if salutation_idx != -1:
        body_start = salutation_idx + 1
        # Save a copy of the first body paragraph element to use as a template
        template_p_element = copy.deepcopy(doc.paragraphs[body_start]._element)
        template_p_element.clear_content()

        # Remove existing body and closing paragraphs
        while len(doc.paragraphs) > body_start:
            p = doc.paragraphs[body_start]
            p._element.getparent().remove(p._element)

    content = """I am writing to express my strong interest.
    
This is the second paragraph.
<verify>Some red text here</verify>"""

    # We append new paragraphs by inserting clones of template_p_element
    parent_element = doc.paragraphs[salutation_idx]._element.getparent()
    
    # We'll just add them to the document
    for line in content.split("\n"):
        if line.strip() and not _is_header_line(line, company_name):
            new_p_element = copy.deepcopy(template_p_element)
            parent_element.append(new_p_element)
            p = Paragraph(new_p_element, doc)
            
            parts = re.split(r'(<verify>.*?</verify>)', line)
            for part in parts:
                if part.startswith("<verify>") and part.endswith("</verify>"):
                    text = part.replace("<verify>", "").replace("</verify>", "")
                    run = p.add_run(text)
                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                else:
                    p.add_run(part)

    # Add back closing paragraphs
    for closing_text in closing_texts:
        if closing_text.strip():
            new_p_element = copy.deepcopy(template_p_element)
            parent_element.append(new_p_element)
            p = Paragraph(new_p_element, doc)
            p.add_run(closing_text)

    doc.save('scratch/cloned_full.docx')
    print("Done")

if __name__ == "__main__":
    test_cloning()
