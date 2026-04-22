from docx import Document

def inspect_docx(path):
    print(f"--- Inspecting {path} ---")
    doc = Document(path)
    print(f"Number of paragraphs: {len(doc.paragraphs)}")
    for i, para in enumerate(doc.paragraphs[:10]):
        print(f"Para {i} style: {para.style.name} | Text: {para.text[:50]}")
    
    # Check sections (margins)
    for i, section in enumerate(doc.sections):
        print(f"Section {i} Margins - Top: {section.top_margin.inches} in, Bottom: {section.bottom_margin.inches} in")

inspect_docx("Aschettino, Ava - Cover Letter Template.docx")
print("\n")
inspect_docx("Aschettino, Ava- Resume.docx")
