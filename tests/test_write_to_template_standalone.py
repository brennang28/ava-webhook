#!/usr/bin/env python3
"""Standalone test for _write_to_template logic without heavy deps."""
import sys
import os
import re
import io
import time

# Add src to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))

# We need to stub out langgraph/langchain before importing generator
class FakeModule:
    pass

# Stub langgraph
sys.modules['langgraph'] = FakeModule()
sys.modules['langgraph.graph'] = FakeModule()
FakeModule.START = 'START'
FakeModule.END = 'END'
FakeModule.StateGraph = lambda *a, **k: FakeModule()
FakeModule.StateGraph.add_node = lambda *a, **k: None
FakeModule.StateGraph.add_edge = lambda *a, **k: None
FakeModule.StateGraph.compile = lambda *a, **k: None

# Stub langchain_ollama
sys.modules['langchain_ollama'] = FakeModule()
sys.modules['langchain_ollama'].ChatOllama = lambda *a, **k: FakeModule()

# Stub langchain_core.messages
sys.modules['langchain_core'] = FakeModule()
sys.modules['langchain_core.messages'] = FakeModule()
sys.modules['langchain_core.messages'].SystemMessage = lambda **k: None
sys.modules['langchain_core.messages'].HumanMessage = lambda **k: None

# Stub google auth
sys.modules['google.oauth2.credentials'] = FakeModule()
sys.modules['google.oauth2.credentials'].Credentials = FakeModule()
sys.modules['google.auth'] = FakeModule()
sys.modules['google.auth.transport'] = FakeModule()
sys.modules['google.auth.transport.requests'] = FakeModule()
sys.modules['google.auth.transport.requests'].Request = lambda *a, **k: None
sys.modules['google.auth.exceptions'] = FakeModule()
sys.modules['google.auth.exceptions'].RefreshError = Exception
sys.modules['googleapiclient.discovery'] = FakeModule()
sys.modules['googleapiclient.discovery'].build = lambda *a, **k: None
sys.modules['googleapiclient.http'] = FakeModule()
sys.modules['googleapiclient.http'].MediaInMemoryUpload = lambda *a, **k: None

# Stub dotenv
sys.modules['dotenv'] = FakeModule()
sys.modules['dotenv'].load_dotenv = lambda: None

# Now we can import the generator
from ava_webhook.generator import AvaGenerator
from docx import Document


def test_preserve_closing_and_replace_placeholders():
    gen = AvaGenerator.__new__(AvaGenerator)
    gen.reasoning_llm = None

    job = {"company": "TestCorp", "role": "Tester", "location": "Remote"}

    content = (
        "Paragraph one about my experience.\n"
        "Paragraph two about my skills.\n"
        "Paragraph three about why TestCorp.\n"
        "Thank you for your time and consideration."
    )

    buffer = gen._write_to_template(
        "assets/Aschettino, Ava - Cover Letter Template.docx",
        content, True, job
    )
    doc = Document(buffer)
    texts = [p.text for p in doc.paragraphs]

    # 1. No placeholders remain
    for t in texts:
        assert "[Company Name]" not in t, f"[Company Name] not replaced in: {t}"
        assert "[Company Address]" not in t, f"[Company Address] not replaced in: {t}"
        assert "[Date]" not in t, f"[Date] not replaced in: {t}"
        assert "[Hiring Manager Name]" not in t, f"[Hiring Manager Name] not replaced in: {t}"

    # 2. Closing preserved
    assert any("Sincerely" in t for t in texts), "Closing 'Sincerely,' missing"
    assert any("Ava Aschettino" in t for t in texts), "Signature missing"

    # 3. Company name in header
    assert any("TestCorp" in t for t in texts), "Company name not in header"

    # 4. Body paragraphs have spacing
    salutation_idx = next((i for i, t in enumerate(texts) if "Dear" in t), -1)
    closing_idx = next((i for i, t in enumerate(texts) if "Sincerely" in t), -1)
    assert salutation_idx != -1, "Salutation not found"
    assert closing_idx != -1, "Closing not found"

    for i in range(salutation_idx + 1, closing_idx):
        para = doc.paragraphs[i]
        assert para.paragraph_format.space_after is not None, \
            f"Body paragraph {i} missing space_after"

    print("PASS: preserve_closing_and_replace_placeholders")


def test_filter_duplicate_headers():
    gen = AvaGenerator.__new__(AvaGenerator)
    gen.reasoning_llm = None

    job = {"company": "Acme Inc", "role": "Dev", "location": "New York, NY"}

    # Simulate LLM output that includes duplicate headers
    content = (
        "Ava Aschettino\n"
        "New York, NY\n"
        "(516) 532-3384\n"
        "avaaschettino@gmail.com\n"
        "Hiring Manager\n"
        "Acme Inc\n"
        "Dear Hiring Manager,\n"
        "This is the real body paragraph.\n"
        "Sincerely,\n"
        "Ava Aschettino\n"
        "Thank you for your time."
    )

    buffer = gen._write_to_template(
        "assets/Aschettino, Ava - Cover Letter Template.docx",
        content, True, job
    )
    doc = Document(buffer)
    texts = [p.text for p in doc.paragraphs]

    salutation_idx = next((i for i, t in enumerate(texts) if "Dear" in t), -1)
    closing_idx = next((i for i, t in enumerate(texts) if "Sincerely" in t), -1)
    assert salutation_idx != -1 and closing_idx != -1

    body_texts = texts[salutation_idx + 1:closing_idx]
    for bt in body_texts:
        # The only body paragraph should be the real one
        if "Ava Aschettino" in bt and "body" not in bt.lower():
            raise AssertionError(f"Duplicate header in body: {bt}")
        if bt.strip() == "Hiring Manager":
            raise AssertionError(f"Duplicate header in body: {bt}")

    print("PASS: filter_duplicate_headers")


def test_resume_spacing():
    gen = AvaGenerator.__new__(AvaGenerator)
    gen.reasoning_llm = None

    job = {"company": "TestCorp", "role": "Tester", "location": "Remote"}
    content = "Line one\nLine two\nLine three"

    buffer = gen._write_to_template(
        "assets/Aschettino, Ava- Resume.docx",
        content, False, job
    )
    doc = Document(buffer)

    assert len(doc.paragraphs) >= 3, f"Expected >=3 paragraphs, got {len(doc.paragraphs)}"

    # Check spacing on paragraphs beyond header
    for p in doc.paragraphs[3:]:
        assert p.paragraph_format.space_after is not None, \
            "Resume paragraph missing space_after"

    print("PASS: resume_spacing")


if __name__ == "__main__":
    test_preserve_closing_and_replace_placeholders()
    test_filter_duplicate_headers()
    test_resume_spacing()
    print("\nAll tests passed!")
