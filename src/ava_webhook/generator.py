import operator
import os
import requests
import json
import re
import time
import io
import logging
from typing import Annotated, TypedDict, List, Dict, Any
from docx import Document
from docx.shared import RGBColor, Pt
from langgraph.graph import StateGraph, START, END
from langchain_ollama import ChatOllama
from langfuse import Langfuse
from langchain_core.messages import SystemMessage, HumanMessage
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from google.auth.exceptions import RefreshError
from googleapiclient.discovery import build
from googleapiclient.http import MediaInMemoryUpload
from dotenv import load_dotenv

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

class OverallState(TypedDict):
    jobs: List[Dict[str, Any]]
    results: Annotated[List[Dict[str, Any]], operator.add]

class JobProcessState(TypedDict):
    job: Dict[str, Any]
    resume_text: str
    template_text: str
    job_analysis: Dict[str, Any]
    mapped_experience: List[Dict[str, Any]]
    final_cover_letter: str
    final_resume_draft: str
    revision_count: int
    critique: str
    accuracy_report: str
    folder_link: str

class AvaGenerator:
    def __init__(self, config_path="config.json", langfuse_client: Langfuse | None = None):
        # 0. Load Config
        self.config = {}
        if os.path.exists(config_path):
            try:
                with open(config_path, 'r') as f:
                    self.config = json.load(f)
            except Exception as e:
                print(f"Error loading config in generator: {e}")
        
        models_config = self.config.get("models", {})
        self.reasoning_model = models_config.get("reasoning", "gemma4:26b")
        self.writing_model = models_config.get("writing", "qwen2.5vl:3b")

        # 1. Models
        self.cloud_url = os.getenv("OLLAMA_CLOUD_URL")
        self.desktop_url = os.getenv("DESKTOP_OLLAMA_URL", "http://127.0.0.1:11434")
        self.ollama_api_key = os.getenv("OLLAMA_AUX2_API_KEY")

        self.reasoning_llm = self._init_llm(self.reasoning_model)
        self.writing_llm = self._init_llm(self.writing_model)

        # 2. Context paths
        self.resume_path = "assets/Aschettino, Ava- Resume.docx"
        self.template_path = "assets/Aschettino, Ava - Cover Letter Template.docx"
        
        # Load text for LLM
        self.resume_text = self._load_context(self.resume_path)
        self.template_text = self._load_context(self.template_path)
        
        # 3. Main Workflow
        self.workflow = self._build_graph()
        self.langfuse = langfuse_client or Langfuse()

    def _init_llm(self, model_name: str, **kwargs) -> ChatOllama:
        """Initializes ChatOllama with appropriate base URL and auth."""
        if ":cloud" in model_name and self.cloud_url:
            return ChatOllama(
                model=model_name,
                base_url=self.cloud_url,
                client_kwargs={"headers": {"Authorization": f"Bearer {self.ollama_api_key}"}} if self.ollama_api_key else {},
                **kwargs
            )
        else:
            return ChatOllama(
                model=model_name,
                base_url=self.desktop_url,
                **kwargs
            )
        
    def _load_context(self, path: str) -> str:
        if not os.path.exists(path):
            return ""
        
        if path.endswith(".docx"):
            try:
                doc = Document(path)
                return "\n".join([p.text for p in doc.paragraphs])
            except Exception as e:
                print(f"Error extracting text from docx: {e}")
                return ""
        else:
            with open(path, 'r') as f:
                return f.read()

    def _build_graph(self) -> StateGraph:
        builder = StateGraph(OverallState)
        builder.add_node("process_jobs", self._process_jobs_sequentially)
        builder.add_node("send_webhooks", self._send_webhooks)
        
        builder.add_edge(START, "process_jobs")
        builder.add_edge("process_jobs", "send_webhooks")
        builder.add_edge("send_webhooks", END)
        return builder.compile()

    def _process_jobs_sequentially(self, state: OverallState):
        with self.langfuse.start_as_current_observation(
            name="process-jobs-sequentially",
            as_type="span",
            input={"job_count": len(state["jobs"])},
        ):
            results = []
            for job in state["jobs"]:
                print(f"--- Processing {job.get('company')} sequentially ---")
                job_state: JobProcessState = {
                "job": job,
                "resume_text": self.resume_text,
                "template_text": self.template_text,
                "revision_count": 0,
                "job_analysis": {},
                "mapped_experience": [],
                "final_cover_letter": "",
                "final_resume_draft": "",
                "critique": "",
                "folder_link": ""
            }
            
            job_state.update(self._analyze_job(job_state))
            job_state.update(self._map_experience(job_state))
            
            while job_state["revision_count"] < 2:
                job_state.update(self._draft_sections(job_state))
                job_state.update(self._verify_accuracy(job_state))
                job_state.update(self._critique_review(job_state))
                if "EXCELLENT" in job_state.get("critique", "").upper():
                    break
            
            job_state.update(self._finalize_job(job_state))
            results.extend(job_state.get("results", []))
            time.sleep(2)
            
        return {"results": results}

    def _parse_json(self, text: str) -> Any:
        clean_text = re.sub(r'```json\s*|\s*```', '', text).strip()
        try:
            return json.loads(clean_text)
        except:
            match = re.search(r'(\{.*\}|\[.*\])', clean_text, re.DOTALL)
            if match:
                try:
                    return json.loads(match.group(0))
                except:
                    pass
        return {"raw": text}

    def _get_llm_with_config(self, llm_instance, config: dict):
        """Re-initializes LLM with parameters from Langfuse config."""
        model_name = config.get("model", llm_instance.model)
        
        params = {}
        if "temperature" in config:
            params["temperature"] = float(config["temperature"])
        if "num_predict" in config:
            params["num_predict"] = int(config["num_predict"])
        if "top_p" in config:
            params["top_p"] = float(config["top_p"])
            
        return self._init_llm(model_name, **params)

    def _with_generation_observation(self, name: str, model: str, prompt: str, model_parameters: dict, invoke_fn):
        with self.langfuse.start_as_current_observation(
            name=name,
            as_type="generation",
            input={
                "prompt_excerpt": prompt[:1500],
                "prompt_length": len(prompt),
            },
            model=model,
            model_parameters={k: v for k, v in model_parameters.items() if v is not None},
        ) as generation:
            result = invoke_fn()
            output = getattr(result, "content", result)
            if isinstance(output, str) and len(output) > 2000:
                output = output[:2000] + "..."
            generation.update(output={"content": output})
            return result

    def _with_tool_observation(self, name: str, input_data: dict):
        return self.langfuse.start_as_current_observation(
            name=name,
            as_type="tool",
            input=input_data,
        )

    def _run_drive_step(self, step_name: str, input_data: dict, action):
        with self._with_tool_observation(
            name=f"upload-to-drive/{step_name}",
            input_data=input_data,
        ) as observation:
            try:
                result = action()
                observation.update(output={"step": step_name, "status": "success"})
                return result
            except Exception as e:
                observation.update(output={"step": step_name, "status": "error", "error": str(e)})
                raise

    def _analyze_job(self, state: JobProcessState):
        job = state["job"]
        print(f"--- [Node: Analyze Job] starting for {job.get('company')} ---")
        
        prompt_obj = self.langfuse.get_prompt("analyze-job", label="production")
        compiled_prompt = prompt_obj.compile(
            role=job.get('role'),
            company=job.get('company'),
            salary=job.get('salary', 'N/A'),
            description=job.get('description', 'No description provided')
        )
        
        llm = self._get_llm_with_config(self.reasoning_llm, prompt_obj.config)
        res = self._with_generation_observation(
            name="analyze-job",
            model=self.reasoning_model,
            prompt=compiled_prompt,
            model_parameters=prompt_obj.config,
            invoke_fn=lambda: llm.invoke([SystemMessage(content="You are a requirement analyst."), HumanMessage(content=compiled_prompt)]),
        )
        analysis = self._parse_json(res.content)
        return {"job_analysis": analysis}

    @staticmethod
    def _normalize_mapped_experience(raw: Any) -> list[dict[str, Any]]:
        if isinstance(raw, list):
            return [item for item in raw if isinstance(item, dict)]
        if isinstance(raw, dict):
            for key in ("mapping", "mappings", "mapped_experience", "experience"):
                val = raw.get(key)
                if isinstance(val, list):
                    return [item for item in val if isinstance(item, dict)]
        return []

    def _map_experience(self, state: JobProcessState):
        analysis = state["job_analysis"]
        resume = state["resume_text"]
        print(f"--- [Node: Map Experience] mapping to STAR evidence ---")
        
        prompt_obj = self.langfuse.get_prompt("map-experience", label="production")
        compiled_prompt = prompt_obj.compile(
            requirements=analysis.get('requirements'),
            resume=resume
        )
        
        llm = self._get_llm_with_config(self.reasoning_llm, prompt_obj.config)
        res = self._with_generation_observation(
            name="map-experience",
            model=self.reasoning_model,
            prompt=compiled_prompt,
            model_parameters=prompt_obj.config,
            invoke_fn=lambda: llm.invoke([SystemMessage(content="You are a career matcher."), HumanMessage(content=compiled_prompt)]),
        )
        mapping = self._parse_json(res.content)
        return {"mapped_experience": self._normalize_mapped_experience(mapping)}

    def _draft_sections(self, state: JobProcessState):
        job = state["job"]
        analysis = state["job_analysis"]
        mapping = state["mapped_experience"]
        template = state["template_text"]
        print(f"--- [Node: Draft Sections] Revision: {state['revision_count']} ---")
        
        # Phase 1: Selection (Resolve Brackets)
        prompt_obj_sel = self.langfuse.get_prompt("resolve-template", label="production")
        selection_prompt = prompt_obj_sel.compile(
            template=template,
            mapping=mapping,
            company=job.get('company')
        )
        
        llm_sel = self._get_llm_with_config(self.reasoning_llm, prompt_obj_sel.config)
        selection_res = self._with_generation_observation(
            name="resolve-template",
            model=self.reasoning_model,
            prompt=selection_prompt,
            model_parameters=prompt_obj_sel.config,
            invoke_fn=lambda: llm_sel.invoke([SystemMessage(content="You are a logic engine that resolves template choices."), HumanMessage(content=selection_prompt)]),
        )
        resolved_template = selection_res.content

        # Phase 2: Writing (Prose Polishing)
        prompt_obj_write = self.langfuse.get_prompt("write-cover-letter", label="production")
        writing_prompt = prompt_obj_write.compile(
            resolved_template=resolved_template,
            role=job.get('role'),
            company=job.get('company'),
            vibe=analysis.get('vibe'),
            mapping=mapping
        )
        
        llm_write = self._get_llm_with_config(self.writing_llm, prompt_obj_write.config)
        res = self._with_generation_observation(
            name="write-cover-letter",
            model=self.writing_model,
            prompt=writing_prompt,
            model_parameters=prompt_obj_write.config,
            invoke_fn=lambda: llm_write.invoke([SystemMessage(content="Persuasive career writer with a focus on brand alignment."), HumanMessage(content=writing_prompt)]),
        )
        return {"final_cover_letter": res.content}

    def _verify_accuracy(self, state: JobProcessState):
        letter = state["final_cover_letter"]
        resume = state["resume_text"]
        job = state["job"]
        print(f"--- [Node: Verify Accuracy] cross-referencing claims ---")
        
        prompt_obj = self.langfuse.get_prompt("verify-accuracy", label="production")
        compiled_prompt = prompt_obj.compile(
            resume=resume,
            letter=letter,
            role=job.get('role'),
            company=job.get('company')
        )
        
        llm = self._get_llm_with_config(self.reasoning_llm, prompt_obj.config)
        res = self._with_generation_observation(
            name="verify-accuracy",
            model=self.reasoning_model,
            prompt=compiled_prompt,
            model_parameters=prompt_obj.config,
            invoke_fn=lambda: llm.invoke([SystemMessage(content="Strict factual auditor."), HumanMessage(content=compiled_prompt)]),
        )
        report = self._parse_json(res.content)
        return {"accuracy_report": json.dumps(report)}

    def _critique_review(self, state: JobProcessState):
        letter = state["final_cover_letter"]
        accuracy = state.get("accuracy_report", "{}")
        print(f"--- [Node: Critique Review] round {state['revision_count']} ---")
        
        prompt_obj = self.langfuse.get_prompt("critique-review", label="production")
        compiled_prompt = prompt_obj.compile(
            accuracy=accuracy,
            letter=letter
        )
        
        llm = self._get_llm_with_config(self.reasoning_llm, prompt_obj.config)
        res = self._with_generation_observation(
            name="critique-review",
            model=self.reasoning_model,
            prompt=compiled_prompt,
            model_parameters=prompt_obj.config,
            invoke_fn=lambda: llm.invoke([SystemMessage(content="Senior hiring editor focused on brand and integrity."), HumanMessage(content=compiled_prompt)]),
        )
        return {"revision_count": state["revision_count"] + 1, "critique": res.content}

    def _finalize_job(self, state: JobProcessState):
        print(f"--- [Node: Finalize Job] generating buffers ---")
        job = state["job"]
        resume_draft = "AVA ASCHETTINO - TAILORED DRAFT\n"
        resume_draft += f"Job: {job.get('role')} at {job.get('company')}\n"
        resume_draft += "="*40 + "\n\n"
        for item in state.get("mapped_experience", []):
            if not isinstance(item, dict):
                continue
            resume_draft += f"REQ: {item.get('requirement')}\nSTAR: {item.get('evidence')}\n\n"
        
        cover_letter_buffer = self._write_to_template(self.template_path, state["final_cover_letter"], True, job)
        resume_draft_buffer = self._write_to_template(self.resume_path, resume_draft, False, job)

        folder_link = self._upload_to_drive(
            job.get("company", "Unknown"), 
            job.get("role", "Unknown"), 
            cover_letter_buffer,
            resume_draft_buffer
        )
        
        processed_job = job.copy()
        processed_job["cover_letter_text"] = state["final_cover_letter"]
        processed_job["resume_draft_text"] = resume_draft
        processed_job["folder_link"] = folder_link
        
        # Use AI-extracted salary if available
        analysis = state.get("job_analysis", {})
        if analysis.get("salary"):
            processed_job["salary"] = analysis["salary"]
            
        return {"results": [processed_job]}

    def _lookup_company_address(self, company_name: str) -> str:
        api_key = os.getenv("TAVILY_API_KEY")
        if not api_key:
            return "New York, NY"
        
        try:
            logger.info(f"Looking up address for {company_name} via Tavily...")
            url = "https://api.tavily.com/search"
            payload = {
                "api_key": api_key,
                "query": f"{company_name} headquarters office address New York",
                "search_depth": "basic",
                "max_results": 1
            }
            response = requests.post(url, json=payload, timeout=10)
            data = response.json()
            
            # Use LLM to extract a clean address from the search results
            context = "\n".join([r.get("content", "") for r in data.get("results", [])])
            if not context:
                return "New York, NY"
                
            prompt_obj = self.langfuse.get_prompt("extract-address", label="production")
            compiled_prompt = prompt_obj.compile(
                company_name=company_name,
                context=context
            )
            
            llm = self._get_llm_with_config(self.reasoning_llm, prompt_obj.config)
            res = self._with_generation_observation(
                name="extract-address",
                model=self.reasoning_model,
                prompt=compiled_prompt,
                model_parameters=prompt_obj.config,
                invoke_fn=lambda: llm.invoke([SystemMessage(content="Address extractor."), HumanMessage(content=compiled_prompt)]),
            )
            return res.content.strip()
        except Exception as e:
            logger.error(f"Address lookup failed: {e}")
            return "New York, NY"

    def _write_to_template(self, template_path: str, content: str, is_cover_letter: bool, job: Dict) -> io.BytesIO:
        doc = Document(template_path)

        if is_cover_letter:
            # 1. Resolve Placeholders
            company_name = job.get("company", "the team")
            job_location = job.get("location", "")

            # Use Tavily only if the job location is missing or a generic city name
            is_generic = job_location.lower() in ["remote", "new york, ny", "nyc", "united states", ""]
            if is_generic:
                address = self._lookup_company_address(company_name)
            else:
                address = job_location

            placeholders = {
                "[Date]": time.strftime("%B %d, %Y"),
                "[Hiring Manager Name]": "Hiring Manager",
                "[Company Name]": company_name,
                "[Company Address]": address
            }

            # Replace placeholders in ALL paragraphs first
            for p in doc.paragraphs:
                for key, val in placeholders.items():
                    if key in p.text:
                        p.text = p.text.replace(key, val)

            # 2. Identify the salutation and closing paragraphs
            salutation_idx = -1
            closing_idx = -1
            for i, p in enumerate(doc.paragraphs):
                if "Dear" in p.text:
                    salutation_idx = i
                if "Sincerely" in p.text:
                    closing_idx = i

            # Strip salutation from LLM content to avoid duplication
            content = re.sub(r'^Dear\s+.*?[,:]\s*', '', content.strip(), flags=re.MULTILINE | re.IGNORECASE)

            # 3. Save closing paragraphs and remove everything after salutation
            closing_texts = []
            if closing_idx != -1 and closing_idx > salutation_idx:
                for i in range(closing_idx, len(doc.paragraphs)):
                    text = doc.paragraphs[i].text
                    if text.strip():
                        closing_texts.append(text)

            if salutation_idx != -1:
                body_start = salutation_idx + 1
                while len(doc.paragraphs) > body_start:
                    p = doc.paragraphs[body_start]
                    p._element.getparent().remove(p._element)

            # 4. Filter header lines from LLM content
            # Use exact-match anchors so legitimate body paragraphs mentioning
            # the company or "Hiring Manager" are not accidentally stripped.
            header_patterns = [
                r"^Ava Aschettino$",
                r"^New York, NY$",
                r"^\(516\) 532-3384$",
                r"^avaaschettino@gmail\.com$",
                r"^Hiring Manager$",
                r"^" + re.escape(company_name) + r"$",
                r"^\[Company Address\]$",
                r"^Sincerely,$",
                r"^Best regards,$",
                r"^Thank you$",
            ]

            def _is_header_line(line: str) -> bool:
                stripped = line.strip()
                for pattern in header_patterns:
                    if re.search(pattern, stripped, re.IGNORECASE):
                        return True
                # Also filter standalone salutations
                if re.match(r'^Dear\s+', stripped, re.IGNORECASE):
                    return True
                return False

            # 5. Append new body content with paragraph spacing
            for line in content.split("\n"):
                if line.strip() and not _is_header_line(line):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(12)
                    # Parse <verify> tags for red-highlighting
                    parts = re.split(r'(<verify>.*?</verify>)', line)
                    for part in parts:
                        if part.startswith("<verify>") and part.endswith("</verify>"):
                            text = part.replace("<verify>", "").replace("</verify>", "")
                            run = p.add_run(text)
                            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                        else:
                            p.add_run(part)

            # 6. Add back closing paragraphs
            for closing_text in closing_texts:
                if closing_text.strip():
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(12)
                    p.add_run(closing_text)
        else:
            header_limit = 3
            while len(doc.paragraphs) > header_limit:
                p = doc.paragraphs[-1]
                p._element.getparent().remove(p._element)

            for line in content.split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(12)
                    # Parse <verify> tags
                    parts = re.split(r'(<verify>.*?</verify>)', line)
                    for part in parts:
                        if part.startswith("<verify>") and part.endswith("</verify>"):
                            text = part.replace("<verify>", "").replace("</verify>", "")
                            run = p.add_run(text)
                            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                        else:
                            p.add_run(part)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _upload_to_drive(self, company: str, role: str, cover_buffer: io.BytesIO, resume_buffer: io.BytesIO) -> str:
        parent_folder_id = "1sonKNR4S-OGhKc2Szg54k2L0EtoOjK1Z"
        token_path = 'token.json'
        backup_dir = "drafts"
        os.makedirs(backup_dir, exist_ok=True)
        sanitized_company = company.replace(' ', '_').replace('/', '_')
        local_cover_path = os.path.join(backup_dir, f"{sanitized_company}_CoverLetter.docx")
        local_resume_path = os.path.join(backup_dir, f"{sanitized_company}_ResumeDraft.docx")
        upload_result = {
            "local_backup": False,
            "drive_upload": False,
            "folder_link": None,
            "error": None,
            "folder_created": False,
            "cover_letter_uploaded": False,
            "resume_uploaded": False,
        }

        with self._with_tool_observation(
            name="upload-to-drive",
            input_data={
                "company": company,
                "role": role,
                "backup_dir": backup_dir,
                "token_path": token_path,
            },
        ) as observation:
            try:
                with open(local_cover_path, "wb") as f:
                    f.write(cover_buffer.getvalue())
                with open(local_resume_path, "wb") as f:
                    f.write(resume_buffer.getvalue())
                upload_result["local_backup"] = True
            except Exception as e:
                upload_result["error"] = f"Local backup failed: {e}"
                logger.error(upload_result["error"])
                observation.update(output=upload_result)
                return f"Local backup failed: {e}"

            if not os.path.exists(token_path):
                upload_result["error"] = "token.json missing"
                observation.update(output=upload_result)
                return f"Saved locally only (token.json missing): {local_cover_path}"

            try:
                creds = Credentials.from_authorized_user_file(token_path, ['https://www.googleapis.com/auth/drive.file'])
                if not creds.valid:
                    if creds.expired and creds.refresh_token:
                        try:
                            creds.refresh(Request())
                            with open(token_path, 'w') as token_file:
                                token_file.write(creds.to_json())
                        except RefreshError as refresh_error:
                            upload_result["error"] = f"Drive auth refresh failed: {refresh_error}"
                            observation.update(output=upload_result)
                            return (
                                "Saved locally! Drive auth failed: token expired or revoked. "
                                "Please rerun scripts/setup_oauth.py to reauthorize."
                            )
                    else:
                        upload_result["error"] = "Drive credentials invalid or expired with no refresh token"
                        observation.update(output=upload_result)
                        return (
                            "Saved locally! Drive auth failed: invalid or expired token. "
                            "Please rerun scripts/setup_oauth.py to reauthorize."
                        )

                drive_service = build('drive', 'v3', credentials=creds)

                folder_metadata = {
                    'name': f"{company} - {role}",
                    'mimeType': 'application/vnd.google-apps.folder',
                    'parents': [parent_folder_id]
                }
                folder = self._run_drive_step(
                    "create-folder",
                    {
                        "company": company,
                        "role": role,
                        "parent_folder_id": parent_folder_id,
                        "mimeType": folder_metadata['mimeType'],
                    },
                    lambda: drive_service.files().create(body=folder_metadata, fields='id, webViewLink').execute(),
                )
                folder_id = folder.get('id')
                upload_result["folder_created"] = True

                cv_metadata = {
                    'name': f'Cover Letter - {company}',
                    'mimeType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'parents': [folder_id]
                }
                cv_media = MediaInMemoryUpload(
                    cover_buffer.getvalue(), 
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                    resumable=True
                )
                self._run_drive_step(
                    "upload-cover-letter",
                    {
                        "company": company,
                        "role": role,
                        "file_name": cv_metadata['name'],
                        "mimeType": cv_metadata['mimeType'],
                        "parent_folder_id": folder_id,
                    },
                    lambda: drive_service.files().create(body=cv_metadata, media_body=cv_media, fields='id').execute(),
                )
                upload_result["cover_letter_uploaded"] = True

                rs_metadata = {
                    'name': f'Resume Tailoring Draft - {company}',
                    'mimeType': 'application/vnd.google-apps.document',
                    'parents': [folder_id]
                }
                rs_media = MediaInMemoryUpload(
                    resume_buffer.getvalue(), 
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                    resumable=True
                )
                self._run_drive_step(
                    "upload-resume",
                    {
                        "company": company,
                        "role": role,
                        "file_name": rs_metadata['name'],
                        "mimeType": rs_metadata['mimeType'],
                        "parent_folder_id": folder_id,
                    },
                    lambda: drive_service.files().create(body=rs_metadata, media_body=rs_media, fields='id').execute(),
                )
                upload_result["resume_uploaded"] = True

                upload_result["drive_upload"] = True
                upload_result["folder_link"] = folder.get('webViewLink', "No link generated")
                observation.update(output=upload_result)
                return upload_result["folder_link"]
            except Exception as e:
                upload_result["error"] = f"Drive upload failed: {e}"
                observation.update(output=upload_result)
                return f"Saved locally! Drive error: {str(e)}"

    def _send_webhooks(self, state: OverallState):
        webhook_url = os.getenv("WEBHOOK_URL")
        if not webhook_url:
             return {}
            
        print(f"--- [Node: Send Webhooks] sending {len(state['results'])} responses ---")
        with self._with_tool_observation(
            name="send-webhooks",
            input_data={"result_count": len(state.get('results', []))},
        ) as observation:
            success_count = 0
            for job in state["results"]:
                try:
                    payload = {
                        "company": job.get("company", "Unknown"),
                        "role": job.get("role") or job.get("title") or "Position",
                        "salary": job.get("salary", "N/A"),
                        "link": job.get("link", ""),
                        "folder_link": job.get("folder_link", "")
                    }
                    response = requests.post(webhook_url, json=payload, timeout=15)
                    if response.status_code == 200:
                        success_count += 1
                    logger.info(f"Dispatched: {payload['role']} @ {payload['company']}")
                except Exception as e:
                    logger.error(f"Error sending webhook: {e}")
            observation.update(output={"success_count": success_count, "payload_count": len(state.get('results', []))})
        return {}
